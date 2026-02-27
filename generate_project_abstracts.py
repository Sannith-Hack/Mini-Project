from docx import Document
from docx.shared import Pt

projects = {
    "Mental_Stress_Detector_Abstract_Final.docx": {
        "title": "Mental Stress Detector",
        "content": [
            "Problem in real life: Modern life and high-pressure college environments often lead to hidden stress. Many students do not realize when their stress levels have become unhealthy until it affects their sleep or health.",
            "Simple solution idea: This project is a digital companion that monitors signs of stress and alerts the user to take a break or practice relaxation before the stress becomes a serious problem.",
            "How the system works: The system uses simple sensors (like a heart rate monitor) or a mobile app to track mood and physical activity. When it detects patterns associated with high stress, it sends a gentle notification suggesting deep breathing or a short walk.",
            "Benefits to society: It promotes mental well-being and helps prevent burnout. By catching stress early, it helps people maintain a healthier lifestyle and prevents long-term health issues.",
            "Why this project is useful: It makes mental health monitoring easy and accessible, providing immediate support to students and professionals in high-pressure environments."
        ]
    },
    "Smart_Attendance_Abstract_Final.docx": {
        "title": "Smart Classroom Attendance using Face Detection",
        "content": [
            "Problem in real life: In many colleges, teachers still call out names one by one to take attendance. This process is slow, eats up 10–15 minutes of every lecture, and allows students to mark 'proxy' attendance for their absent friends.",
            "Simple solution idea: The solution is an automated system that uses a camera to 'recognize' students as they enter or sit in the classroom. It replaces the manual register with a digital eye that knows exactly who is present.",
            "How the system works: A camera is mounted at the front of the class. It captures a quick photo of the room and compares the faces against a database of registered student photos. If a match is found, the student is marked 'Present' in a digital sheet automatically.",
            "Benefits to society: It saves valuable teaching time and ensures that attendance records are 100% honest. It also helps parents and authorities get real-time updates on student presence without any paperwork.",
            "Why this project is useful: This project brings modern automation to education, making the classroom more disciplined and efficient while removing the possibility of human error or cheating in attendance."
        ]
    },
    "Smart_Medicine_Reminder_Abstract_Final.docx": {
        "title": "Smart Medicine Reminder & Health Tracker",
        "content": [
            "Problem in real life: Elderly people or patients with long-term illnesses often have to take multiple medicines at different times of the day. Forgetting even one dose can lead to serious health complications or emergency hospital visits.",
            "Simple solution idea: This project is a smart medicine box or app that acts as a dedicated health companion. It doesn't just store medicines; it actively reminds the patient when it is time to take a specific pill.",
            "How the system works: The user or a caregiver sets the medication schedule into the system. At the scheduled time, the device makes a sound or sends a notification. It can also track if the medicine was actually taken and send an alert to a family member if a dose is missed.",
            "Benefits to society: It provides peace of mind to families and independence to the elderly. It reduces the risk of medical errors and helps people recover faster by ensuring they follow their doctor's orders strictly.",
            "Why this project is useful: This project saves lives by preventing accidental missed doses and healthcare management simple and stress-free for those who need it most."
        ]
    },
    "Smart_Watering_Advisor_Abstract_Final.docx": {
        "title": "Smart Crop Watering Advisor for Farmers",
        "content": [
            "Problem in real life: Farmers often guess when to water their crops. If they water too much, they waste precious water and can rot the roots; if they water too little, the crops dry up and die. This uncertainty leads to poor harvests and financial loss.",
            "Simple solution idea: The idea is a 'smart assistant' for the field that tells the farmer exactly when the soil is thirsty. It acts like a health monitor for the farm, ensuring plants get water only when they actually need it.",
            "How the system works: Small sensors are placed in the soil to measure moisture levels. These sensors send data to a central unit that also checks the local weather forecast. If the soil is dry and no rain is expected, the system alerts the farmer or automatically turns on the water pump.",
            "Benefits to society: It conserves a huge amount of water and helps farmers grow healthier crops with less effort. This increases food production and supports the livelihoods of small-scale farmers.",
            "Why this project is useful: By using data instead of guesswork, this project makes farming more scientific and sustainable, helping to solve the global problem of water scarcity in agriculture."
        ]
    },
    "Custom_Windows_ISO_Abstract_Final.docx": {
        "title": "Custom Windows ISO for B.Tech Studies",
        "content": [
            "Problem in real life: Whenever a computer science student buys a new laptop or resets their PC, they spend an entire day downloading and installing software. Finding the right versions of Python, Java, and Microsoft Office, and setting up complex tools like WSL, is frustrating and time-consuming.",
            "Simple solution idea: The solution is a 'Pre-Configured Student Windows Kit.' Instead of a blank Windows installation, this is a customized version that comes 'ready-to-code' out of the box, with every required tool already installed and configured.",
            "How the system works: The project involves taking a standard Windows image and 'injecting' essential software like Python, Java, VS Code, and WSL into it using automation scripts. When a student installs Windows using this custom file, all their lab tools are already there on the desktop the first time they log in.",
            "Benefits to society: It saves thousands of hours for students and college lab technicians. It ensures that every student in a class has the exact same software setup, which eliminates 'installation errors' during important lab sessions or exams.",
            "Why this project is useful: This project removes the technical headache of software setup, allowing students to focus immediately on learning and building projects rather than troubleshooting their computers."
        ]
    }
}

team_members = [
    {"name": "P.Sannith (Leader)", "roll": "23567T0942"},
    {"name": "R.Jayathusri", "roll": "23567T0951"},
    {"name": "K.Akshitha", "roll": "23567T0931"}
]

branch = "CSE"
year = "3rd Year"

def create_docx(filename, data):
    doc = Document()
    doc.add_heading(data["title"], 0)
    
    # Abstract Content
    doc.add_heading("ABSTRACT", level=1)
    for para in data["content"]:
        p = doc.add_paragraph()
        if ": " in para:
            parts = para.split(": ", 1)
            run = p.add_run(parts[0] + ": ")
            run.bold = True
            p.add_run(parts[1])
        else:
            p.add_run(para)
            
    doc.add_paragraph() # Spacer
    
    # Team Details at the Bottom
    doc.add_heading("TEAM DETAILS", level=2)
    
    p = doc.add_paragraph()
    p.add_run("Branch: " + branch + "\n").bold = True
    p.add_run("Year:   " + year).bold = True

    # Team Table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "S.No"
    hdr_cells[1].text = "Student Name"
    hdr_cells[2].text = "Roll Number"
    
    # Add members
    for i, member in enumerate(team_members, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = member["name"]
        row_cells[2].text = member["roll"]
            
    doc.save(filename)

if __name__ == "__main__":
    for filename, data in projects.items():
        create_docx(filename, data)
        print("Finalized " + filename)
